import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from agents.scorer import ApplicationDecision
from automation.package_builder import build_application_package
from automation.tailoring import TailoredDocuments
from core.models import Job


class PackageBuilderTests(unittest.TestCase):
    def test_package_contains_employer_files_but_not_internal_intelligence(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            cv = root / "tailored_cv.docx"
            cover = root / "tailored_cover_letter.docx"
            cert = root / "certificate.pdf"
            for path, text in ((cv, "Executive support — calendar management"), (cover, "Dear Hiring Team")):
                doc = Document(); doc.add_paragraph(text); doc.save(path)
            cert.write_bytes(b"%PDF-1.4 supporting document")
            documents = TailoredDocuments("EXECUTIVE_OPERATIONS", root, cv, cover, cert, ("executive support",))
            decision = ApplicationDecision(82, 80, 86, True, ("executive support",), ("salesforce",), "Eligible for automatic application")
            job = Job("Example Ltd", "Executive_Assistant", "Remote", "test", "https://example.test/job")

            package = build_application_package(job, documents, decision)
            self.assertTrue(package.cv_pdf.is_file())
            self.assertTrue(package.cover_letter_pdf.is_file())
            self.assertTrue(package.internal_manifest.is_file())
            self.assertTrue(package.archive.is_file())
            self.assertGreater(package.cv_pdf.stat().st_size, 100)
            self.assertEqual(package.cv_pdf.read_bytes()[:4], b"%PDF")
            self.assertNotIn("_", package.cv_pdf.name)
            self.assertNotIn("_", package.cover_letter_pdf.name)
            self.assertNotIn("—", package.cv_pdf.name)
            self.assertNotIn("–", package.cv_pdf.name)

            manifest = package.internal_manifest.read_text(encoding="utf-8")
            self.assertIn("salesforce", manifest)
            with zipfile.ZipFile(package.archive) as bundle:
                names = bundle.namelist()
                self.assertIn(package.cv_pdf.name, names)
                self.assertIn(package.cover_letter_pdf.name, names)
                self.assertIn(cert.name, names)
                self.assertNotIn(package.internal_manifest.name, names)
                self.assertFalse(any("INTERNAL" in name for name in names))
                self.assertFalse(any("_" in name for name in names if name.lower().endswith((".pdf", ".docx"))))

    def test_repeat_build_replaces_package_cleanly(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            cv = root / "cv.docx"; cover = root / "cover.docx"
            for path in (cv, cover):
                doc = Document(); doc.add_paragraph("Application text"); doc.save(path)
            documents = TailoredDocuments("EXECUTIVE_OPERATIONS", root, cv, cover, None, ())
            decision = ApplicationDecision(90, 90, 90, True, (), (), "Eligible")
            job = Job("Example", "Executive Assistant", "Remote", "test", "https://example.test/job")
            first = build_application_package(job, documents, decision)
            stale = first.folder / "stale.txt"; stale.write_text("remove me")
            second = build_application_package(job, documents, decision)
            self.assertFalse(stale.exists())
            self.assertTrue(second.archive.is_file())


if __name__ == "__main__":
    unittest.main()
