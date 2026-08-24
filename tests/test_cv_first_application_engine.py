import tempfile
from dataclasses import replace
from pathlib import Path

from docx import Document

from automation.candidate_knowledge import load_candidate_knowledge
from automation.profile import ApplicantProfile, EmploymentRecord, remember_verified_answer
from automation.semantic_answers import answer_application_question
from core.text_rules import has_forbidden_dashes, has_forbidden_filename_separator, sanitize_user_filename


def profile_with_cv(root: Path) -> ApplicantProfile:
    cv = root / "Master CV.docx"
    doc = Document()
    doc.add_heading("Professional Experience", level=1)
    doc.add_paragraph("Executive and administrative support. Managed calendars, meetings, executive correspondence, stakeholder coordination and detailed records.")
    doc.add_paragraph("Legal Officer. Prepared complex legal documents and case files with high accuracy, managed client communication and case schedules.")
    doc.add_paragraph("Human Resources Manager. Managed onboarding, personnel records, internal communications and administrative reporting.")
    doc.save(cv)
    cover = root / "cover.docx"; Document().save(cover)
    cert = root / "cert.pdf"; cert.write_bytes(b"pdf")
    return ApplicantProfile(
        first_name="Test", middle_name="Candidate", last_name="User", email="test@example.com",
        phone="+234000000000", city="Abuja", country="Nigeria", address="Abuja", postal_code="900001",
        linkedin_url="", website_url="", work_authorization="Requires sponsorship", requires_sponsorship=True,
        notice_period="Immediately available", salary_expectation="", nationality="Nigerian", region="Africa",
        number_of_employers="4", travel_commitment="Yes", resume_path=str(cv), cover_letter_path=str(cover),
        supporting_document_path=str(cert), employment_history=(
            EmploymentRecord("Executive and administrative support", "Example Admin", "2021", "2022", summary="Managed calendars, meetings, correspondence and stakeholders."),
            EmploymentRecord("Legal Officer", "Example Legal", "2022", "2023", summary="Prepared legal documents, case files and schedules with high accuracy."),
            EmploymentRecord("Human Resources Manager", "Example HR", "2024", current=True, summary="Managed onboarding, records, communications and reporting."),
        ),
    )


def test_master_cv_is_primary_semantic_evidence():
    with tempfile.TemporaryDirectory() as directory:
        p = profile_with_cv(Path(directory))
        knowledge = load_candidate_knowledge(p)
        ranked = knowledge.rank("administration executive support scheduling stakeholder management", "Executive Assistant")
        assert ranked
        assert ranked[0].source.startswith("cv.")
        assert "calendars" in ranked[0].text.lower()


def test_canonical_admin_question_gets_substantive_cv_answer():
    with tempfile.TemporaryDirectory() as directory:
        p = profile_with_cv(Path(directory))
        answer = answer_application_question(
            "Describe any experience with administration, executive support, logistics, events or travel management",
            p,
            job_context="Executive Assistant Corporate Administrator scheduling meetings stakeholders",
        )
        assert answer is not None
        assert len(answer.value) > 80
        assert "calendar" in answer.value.lower() or "administrative" in answer.value.lower()
        assert answer.value.lower() != "yes"


def test_canonical_accuracy_question_uses_cv_evidence():
    with tempfile.TemporaryDirectory() as directory:
        p = profile_with_cv(Path(directory))
        answer = answer_application_question(
            "Describe an occasion where you needed to work through a large amount of detail with high accuracy",
            p,
            job_context="Executive Assistant Corporate Administrator",
        )
        assert answer is not None
        assert "legal" in answer.value.lower() or "case" in answer.value.lower() or "document" in answer.value.lower()


def test_canonical_structured_facts_resolve_without_invention():
    with tempfile.TemporaryDirectory() as directory:
        p = profile_with_cv(Path(directory))
        assert answer_application_question("Which region are you based in?", p).value == "Africa"
        assert answer_application_question("In which country do you currently work?", p).value == "Nigeria"
        assert answer_application_question("In the past ten years, how many companies have you worked for?", p).value == "4"
        assert answer_application_question("Are you willing and able to commit to international travel?", p).value == "Yes"


def test_unknown_school_performance_is_not_invented():
    with tempfile.TemporaryDirectory() as directory:
        p = profile_with_cv(Path(directory))
        assert answer_application_question("How did you perform in mathematics at high school?", p) is None
        assert answer_application_question("Please provide evidence for your high school performance", p) is None


def test_verified_candidate_input_can_be_remembered():
    with tempfile.TemporaryDirectory() as directory:
        root = Path(directory); p = profile_with_cv(root); profile_path = root / "profile.json"
        updated = remember_verified_answer(p, "How did you perform in mathematics at high school?", "Candidate supplied result", profile_path)
        assert updated.standard_answers["how did you perform in mathematics at high school?"] == "Candidate supplied result"
        assert profile_path.is_file()


def test_output_invariants_remove_typographic_dashes_and_filename_underscores():
    assert not has_forbidden_dashes("Clean applicant text")
    name = sanitize_user_filename("Canonical_Executive—Assistant_CV")
    assert name == "Canonical Executive, Assistant CV"
    assert not has_forbidden_filename_separator(name)
