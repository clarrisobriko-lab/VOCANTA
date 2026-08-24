from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document

from automation.profile import ApplicantProfile
from core.text_rules import sanitize_applicant_text


STOP_WORDS = {
    "the", "and", "for", "with", "your", "you", "our", "this", "that", "are", "was",
    "have", "has", "had", "job", "role", "work", "experience", "describe", "tell", "about",
    "please", "candidate", "position", "what", "how", "where", "when", "which", "from",
}


@dataclass(frozen=True, slots=True)
class EvidenceChunk:
    text: str
    source: str
    score: float = 0.0


@dataclass(frozen=True, slots=True)
class CandidateKnowledge:
    chunks: tuple[EvidenceChunk, ...]

    def rank(self, question: str, job_context: str = "", *, limit: int = 6) -> tuple[EvidenceChunk, ...]:
        q = _tokens(question)
        j = _tokens(job_context)
        ranked: list[EvidenceChunk] = []
        for chunk in self.chunks:
            tokens = _tokens(chunk.text)
            if not tokens:
                continue
            direct = len(q & tokens)
            vacancy = len(j & tokens)
            phrase_bonus = sum(2 for phrase in _phrases(question) if phrase in chunk.text.lower())
            score = direct * 4 + vacancy * 1.5 + phrase_bonus
            if score > 0:
                ranked.append(EvidenceChunk(chunk.text, chunk.source, score))
        ranked.sort(key=lambda item: item.score, reverse=True)
        return tuple(ranked[:limit])

    def narrative(self, question: str, job_context: str = "", *, max_chars: int = 1400) -> str:
        ranked = self.rank(question, job_context, limit=5)
        if not ranked:
            return ""
        sentences: list[str] = []
        seen: set[str] = set()
        for item in ranked:
            for sentence in re.split(r"(?<=[.!?])\s+", item.text):
                cleaned = sanitize_applicant_text(sentence).strip(" .")
                key = cleaned.lower()
                if not cleaned or key in seen:
                    continue
                seen.add(key)
                if cleaned[0].islower():
                    cleaned = cleaned[0].upper() + cleaned[1:]
                sentences.append(cleaned.rstrip(".") + ".")
                if len(" ".join(sentences)) >= max_chars:
                    break
            if len(" ".join(sentences)) >= max_chars:
                break
        return " ".join(sentences)[:max_chars].strip()


def _tokens(value: str) -> set[str]:
    return {
        token for token in re.findall(r"[a-z0-9]+", (value or "").lower())
        if len(token) > 2 and token not in STOP_WORDS
    }


def _phrases(value: str) -> tuple[str, ...]:
    text = " ".join((value or "").lower().split())
    phrases = (
        "executive support", "calendar management", "administrative support", "stakeholder management",
        "travel management", "high accuracy", "employee relations", "legal research", "records management",
        "project coordination", "client communication", "internal communications", "case management",
    )
    return tuple(p for p in phrases if p in text)


def _docx_chunks(path: Path) -> list[EvidenceChunk]:
    if not path.is_file() or path.suffix.lower() != ".docx":
        return []
    try:
        document = Document(path)
    except Exception:
        return []
    chunks: list[EvidenceChunk] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = sanitize_applicant_text(paragraph.text)
        if len(text) >= 20:
            chunks.append(EvidenceChunk(text, f"cv.paragraph.{index}"))
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            text = sanitize_applicant_text(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
            if len(text) >= 20:
                chunks.append(EvidenceChunk(text, f"cv.table.{table_index}.{row_index}"))
    return chunks


def _profile_chunks(profile: ApplicantProfile) -> list[EvidenceChunk]:
    chunks: list[EvidenceChunk] = []
    education = profile.highest_education
    chunks.append(EvidenceChunk(
        sanitize_applicant_text(f"{education.degree}, {education.discipline}, {education.institution}, {education.graduation_year}, {education.country}"),
        "profile.education",
    ))
    for index, item in enumerate(profile.employment_history):
        dates = " to ".join(part for part in (item.start_year, "Present" if item.current else item.end_year) if part)
        text = ", ".join(part for part in (item.title, item.employer, dates) if part)
        if item.summary:
            text += f". {item.summary}"
        if text:
            chunks.append(EvidenceChunk(sanitize_applicant_text(text), f"profile.employment.{index}"))
    for question, answer in profile.standard_answers.items():
        if answer:
            chunks.append(EvidenceChunk(sanitize_applicant_text(answer), f"profile.standard_answers.{question}"))
    return chunks


def load_candidate_knowledge(profile: ApplicantProfile) -> CandidateKnowledge:
    cv_chunks = _docx_chunks(Path(profile.resume_path).expanduser())
    profile_chunks = _profile_chunks(profile)
    # CV is primary. Profile facts supplement the CV only where the CV lacks detail.
    return CandidateKnowledge(tuple(cv_chunks + profile_chunks))
