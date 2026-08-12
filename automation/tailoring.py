import re
from dataclasses import dataclass, replace
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from automation.profile import ApplicantProfile
from config.settings import TAILORED_APPLICATIONS_DIR, TAILORING_MAX_KEYWORDS
from core.models import Job


@dataclass(frozen=True, slots=True)
class TailoredDocuments:
    category: str
    folder: Path
    resume_path: Path
    cover_letter_path: Path
    certificate_path: Path | None
    matched_keywords: tuple[str, ...]


CATEGORY_TERMS = {
    "EXECUTIVE_OPERATIONS": ("executive assistant", "administrative assistant", "operations", "project coordinator", "office manager", "virtual assistant", "business support", "administrator", "programme coordinator"),
    "HR_PEOPLE": ("human resources", "hr ", "recruitment", "talent acquisition", "people operations", "employee relations", "onboarding"),
    "LEGAL_COMPLIANCE": ("legal", "counsel", "paralegal", "compliance", "regulatory", "contracts", "governance", "policy", "risk"),
    "NGO_PROGRAMME": ("ngo", "charity", "foundation", "humanitarian", "human rights", "programme officer", "nonprofit", "non-profit", "refugee"),
}

CATEGORY_HEADLINES = {
    "EXECUTIVE_OPERATIONS": "EXECUTIVE OPERATIONS AND ADMINISTRATIVE SPECIALIST",
    "HR_PEOPLE": "HUMAN RESOURCES AND PEOPLE OPERATIONS PROFESSIONAL",
    "LEGAL_COMPLIANCE": "LEGAL, COMPLIANCE AND OPERATIONS PROFESSIONAL",
    "NGO_PROGRAMME": "PROGRAMME, OPERATIONS AND NONPROFIT SUPPORT PROFESSIONAL",
}

SKILL_KEYWORDS = (
    "executive support", "calendar management", "email management", "stakeholder management", "project coordination", "workflow management", "administrative support", "operations", "recruitment", "onboarding", "employee relations", "human resources", "people operations", "documentation", "records management", "client communication", "compliance", "legal research", "contract management", "policy", "google workspace", "microsoft office", "slack", "zoom", "confidential information", "reporting", "scheduling", "training",
)

CATEGORY_SUMMARIES = {
    "EXECUTIVE_OPERATIONS": "Executive Operations and Administrative Specialist with experience supporting leadership, coordinating workflows and managing people, documents and stakeholder communication across corporate, legal and nonprofit environments.",
    "HR_PEOPLE": "Human Resources and People Operations professional with experience in recruitment, onboarding, employee relations, policy implementation, records management and administrative reporting.",
    "LEGAL_COMPLIANCE": "Legal and Compliance professional with experience in legal operations, client management, documentation, research, case coordination, policy support and confidential records administration.",
    "NGO_PROGRAMME": "Operations and Programme Support professional with legal, human resources and nonprofit experience, including stakeholder engagement, human rights advocacy, documentation and administrative coordination.",
}

BASE_SKILLS = {
    "EXECUTIVE_OPERATIONS": ["Executive support and stakeholder management", "Calendar and email management", "Project and workflow coordination", "Administrative reporting and record keeping", "Client communication and relationship management", "Meeting scheduling and travel coordination", "Microsoft Office and Google Workspace", "Remote collaboration using Zoom and Slack", "Recruitment and onboarding", "Legal and compliance administration"],
    "HR_PEOPLE": ["Recruitment and onboarding", "HR operations and employee relations", "Policy implementation and compliance", "Personnel records management", "Training coordination", "Workforce planning and administrative reporting", "Stakeholder communication", "Calendar and workflow management", "Microsoft Office and Google Workspace", "Confidential information management"],
    "LEGAL_COMPLIANCE": ["Legal documentation and case administration", "Compliance and policy support", "Client communication and consultation scheduling", "Legal research and drafting", "Contract and records management", "Stakeholder coordination", "Confidential information management", "Administrative reporting", "Microsoft Office and Google Workspace", "Project and workflow coordination"],
    "NGO_PROGRAMME": ["Programme and administrative coordination", "Stakeholder and government engagement", "Human rights and access to justice support", "Project and workflow management", "Documentation and records management", "Client and beneficiary communication", "Recruitment and onboarding", "Compliance administration", "Microsoft Office and Google Workspace", "Remote team collaboration"],
}

VERIFIED_EXPERIENCE = [
    ("Human Resource Manager, Malachy Godian Enterprises", "January 2025 - Present", ["Oversee recruitment, onboarding and employee management processes.", "Develop and implement policies to improve operational efficiency and compliance.", "Manage employee relations, internal communications and workforce planning.", "Maintain personnel records and provide administrative reports to leadership."]),
    ("Human Resource Manager, Jam Oil and Gas Limited", "January 2024 - December 2024", ["Coordinated onboarding, training programmes and staff scheduling.", "Maintained personnel records and supported management reporting.", "Handled internal communications across multiple departments."]),
    ("Legal Officer and Associate Counsel, Legal Defence and Assistance Project", "July 2022 - December 2023", ["Managed client communication, consultations and case schedules.", "Prepared legal documents, briefs and case files.", "Represented more than 30 clients in civil and criminal matters.", "Worked with NGOs and government agencies on human rights initiatives."]),
    ("Legal Associate, Malcolm Omirhobo & Co.", "February 2021 - June 2022", ["Drafted contracts, motions and legal correspondence.", "Conducted legal research and organised case documentation.", "Managed administrative processes for multiple legal matters."]),
    ("HR Personnel, Micura Services Limited", "January 2019 - February 2021", ["Supported recruitment and employee record management.", "Digitised filing systems and improved document accessibility.", "Managed calendars, meetings and executive correspondence."]),
]

EXPERIENCE_PRIORITY_TERMS = {
    "EXECUTIVE_OPERATIONS": ("executive", "calendar", "meeting", "scheduling", "administrative", "reporting", "communication", "coordination"),
    "HR_PEOPLE": ("human resource", "hr personnel", "recruitment", "onboarding", "employee", "workforce", "personnel", "training"),
    "LEGAL_COMPLIANCE": ("legal", "counsel", "contracts", "case", "compliance", "policy", "research", "client"),
    "NGO_PROGRAMME": ("ngo", "human rights", "government", "legal defence", "access to justice", "programme", "stakeholder", "client"),
}

EXPERIENCE_TITLE_TERMS = {
    "EXECUTIVE_OPERATIONS": ("executive", "administrative", "hr personnel"),
    "HR_PEOPLE": ("human resource", "hr personnel"),
    "LEGAL_COMPLIANCE": ("legal", "counsel"),
    "NGO_PROGRAMME": ("legal defence", "programme", "ngo"),
}


def safe_name(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    return cleaned.strip("_")[:80] or "application"


def classify_job(job: Job) -> str:
    text = f"{job.company} {job.title} {job.description}".lower()
    scores = {category: sum(1 for term in terms if term in text) for category, terms in CATEGORY_TERMS.items()}
    return max(scores, key=lambda category: (scores[category], category == "EXECUTIVE_OPERATIONS"))


def extract_keywords(job: Job) -> tuple[str, ...]:
    text = f"{job.title} {job.description}".lower()
    return tuple(keyword for keyword in SKILL_KEYWORDS if keyword in text)[:TAILORING_MAX_KEYWORDS]


def prioritize_experience(category: str):
    """Rank verified experience by role relevance, with titles carrying extra weight."""
    terms = EXPERIENCE_PRIORITY_TERMS.get(category, ())
    title_terms = EXPERIENCE_TITLE_TERMS.get(category, ())

    def relevance(item) -> int:
        title, _, bullets = item
        title_text = title.lower()
        body_text = " ".join(bullets).lower()
        title_score = sum(3 for term in title_terms if term in title_text)
        body_score = sum(body_text.count(term) for term in terms)
        general_title_score = sum(title_text.count(term) for term in terms)
        return title_score + general_title_score + body_score

    ranked = sorted(enumerate(VERIFIED_EXPERIENCE), key=lambda pair: (-relevance(pair[1]), pair[0]))
    return tuple(item for _, item in ranked)


def _add_contact(document: Document, profile: ApplicantProfile) -> None:
    heading = document.add_heading(profile.full_name.upper(), level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact = document.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact.add_run(f"{profile.city}, {profile.country} | {profile.phone} | {profile.email}\nLinkedIn: {profile.linkedin_url}")


def build_tailored_cv(job: Job, profile: ApplicantProfile, output: Path) -> None:
    category = classify_job(job)
    keywords = extract_keywords(job)
    document = Document()
    _add_contact(document, profile)
    document.add_heading(CATEGORY_HEADLINES[category], level=1)
    document.add_heading("Professional Profile", level=1)
    summary = CATEGORY_SUMMARIES[category]
    if keywords:
        summary += " Relevant strengths include " + ", ".join(keywords[:6]) + "."
    document.add_paragraph(summary)
    document.add_heading("Core Competencies", level=1)
    skills = list(BASE_SKILLS[category])
    for keyword in keywords:
        formatted = keyword.title()
        if formatted.lower() not in {skill.lower() for skill in skills}:
            skills.insert(0, formatted)
    for skill in skills[:12]:
        document.add_paragraph(skill, style="List Bullet")
    document.add_heading("Professional Experience", level=1)
    for title, dates, bullets in prioritize_experience(category):
        document.add_heading(title, level=2)
        document.add_paragraph(dates)
        for bullet in bullets:
            document.add_paragraph(bullet, style="List Bullet")
    document.add_heading("Leadership and Memberships", level=1)
    for item in ("Director of International Engagements, Street Kid Africa Foundation", "Nigerian Bar Association", "African Bar Association", "Commonwealth Lawyers Association"):
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("Education and Professional Development", level=1)
    for item in ("High Impact Executive Assistant Training Program, Skill2Scale Digital, 2026", "Barrister at Law, Nigerian Law School, 2019 - 2020", "Bachelor of Laws, University of Uyo, 2013 - 2018", "Advanced Diploma in Aviation Management, College of Aviation Studies, 2008 - 2010", "Diploma in Community Development, University of Benin, 2005 - 2007"):
        document.add_paragraph(item, style="List Bullet")
    document.save(output)


def build_tailored_cover_letter(job: Job, profile: ApplicantProfile, output: Path) -> None:
    category = classify_job(job)
    keywords = extract_keywords(job)
    document = Document()
    _add_contact(document, profile)
    document.add_paragraph(date.today().strftime("%d %B %Y"))
    document.add_paragraph("Dear Hiring Team,")
    document.add_paragraph(f"I am writing to apply for the {job.title} position at {job.company}. {CATEGORY_SUMMARIES[category]}")
    relevant = ", ".join(keywords[:5]) if keywords else "stakeholder management, workflow coordination and confidential administration"
    document.add_paragraph(f"My experience aligns with the role's emphasis on {relevant}. In my current human resources leadership role, I coordinate recruitment, onboarding, employee relations, records and administrative reporting. My earlier legal and nonprofit work strengthened my client communication, documentation, scheduling and cross-functional coordination capabilities.")
    document.add_paragraph("I bring a practical combination of executive support, operations, people management and compliance awareness. I am comfortable working independently, handling sensitive information and maintaining dependable communication across remote and international teams.")
    document.add_paragraph(f"I would welcome the opportunity to discuss how I can support {job.company} and contribute to the successful delivery of this role. Thank you for your time and consideration.")
    document.add_paragraph(f"Kind regards,\n\n{profile.full_name}")
    document.save(output)


def tailor_documents(job: Job, job_id: int, profile: ApplicantProfile) -> TailoredDocuments:
    category = classify_job(job)
    folder = TAILORED_APPLICATIONS_DIR / f"{job_id}_{safe_name(job.company)}_{safe_name(job.title)}"
    folder.mkdir(parents=True, exist_ok=True)
    resume = folder / "tailored_cv.docx"
    cover = folder / "tailored_cover_letter.docx"
    build_tailored_cv(job, profile, resume)
    build_tailored_cover_letter(job, profile, cover)
    certificate = None
    if category in {"EXECUTIVE_OPERATIONS", "NGO_PROGRAMME"}:
        source = Path(profile.supporting_document_path)
        if source.is_file():
            certificate = folder / source.name
            if source.resolve() != certificate.resolve():
                certificate.write_bytes(source.read_bytes())
    return TailoredDocuments(category=category, folder=folder, resume_path=resume, cover_letter_path=cover, certificate_path=certificate, matched_keywords=extract_keywords(job))


def with_tailored_documents(profile: ApplicantProfile, documents: TailoredDocuments) -> ApplicantProfile:
    return replace(profile, resume_path=str(documents.resume_path), cover_letter_path=str(documents.cover_letter_path), supporting_document_path=(str(documents.certificate_path) if documents.certificate_path else ""))